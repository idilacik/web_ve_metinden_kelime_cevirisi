from flask import Flask, render_template, request, send_file
from ceviriApp import translate_from_url, translate_from_text
from docx import Document
from io import BytesIO
import json
from datetime import datetime
import nltk
import ssl


try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

# NLTK verilerini kontrol et ve indir
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger')

try:
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('wordnet')


app = Flask(__name__)

@app.route('/')
def index():
    return render_template("index.html")
@app.route('/result', methods=['POST'])
def result():

    girilen_url = request.form.get('girilen')
    girilen_metin = request.form.get('metin')

    if girilen_url:
        word_dict = translate_from_url(girilen_url)
        return render_template("result.html", kelimeler=word_dict, url=girilen_url)
    elif girilen_metin:
        word_dict = translate_from_text(girilen_metin)

    return render_template("result.html", kelimeler=word_dict, metin =girilen_metin)


@app.route('/download_word', methods=['POST'])
def download_word():
    kelimeler = request.form.get('kelimeler')


    import json
    kelime_dict = json.loads(kelimeler)

    doc = Document()
    doc.add_heading('İngilizce - Türkçe Kelime Listesi', 0)

    doc.add_paragraph('')  # Boşluk

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'İngilizce'
    hdr_cells[1].text = 'Türkçe'

    for eng, tr in sorted(kelime_dict.items()):
        row_cells = table.add_row().cells
        row_cells[0].text = eng
        row_cells[1].text = tr

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f'kelime_listesi_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    app.run(debug=True, port=5000)